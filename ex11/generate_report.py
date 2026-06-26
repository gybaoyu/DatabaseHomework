"""
生成实验十一 实验报告 Word 文档
"""
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime


def create_report():
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)

    # 标题
    title = doc.add_heading('实验十一 数据库设计实验报告', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('')
    info = [
        '实验名称：数据库设计',
        f'实验日期：{datetime.date.today()}',
        '数据库系统：MySQL 8.0',
        '开发语言：Python 3 + pymysql',
    ]
    for line in info:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # 一、实验目的
    doc.add_heading('一、实验目的', level=1)
    doc.add_paragraph('1. 掌握数据库设计的基本步骤：需求分析、概念结构设计、逻辑结构设计、物理结构设计。')
    doc.add_paragraph('2. 掌握E-R图的绘制方法，理解实体、属性和联系的概念。')
    doc.add_paragraph('3. 掌握E-R图向关系模型的转换方法。')
    doc.add_paragraph('4. 掌握使用SQL语句在MySQL中创建数据库和表的方法。')
    doc.add_paragraph('5. 掌握使用Python（pymysql）连接MySQL数据库并进行增删改查操作。')

    # 二、需求分析
    doc.add_heading('二、需求分析', level=1)
    doc.add_paragraph(
        '本实验设计一个简单的"学生选课管理系统"，实现以下功能：\n'
        '1. 学生信息管理：存储学生的基本信息（学号、姓名、性别、出生日期、所在院系）。\n'
        '2. 课程信息管理：存储课程的基本信息（课程号、课程名、学分、先修课程）。\n'
        '3. 选课信息管理：记录学生的选课情况和考试成绩。'
    )

    # 三、概念结构设计
    doc.add_heading('三、概念结构设计（E-R图）', level=1)

    doc.add_heading('3.1 实体分析', level=2)
    doc.add_paragraph('（1）学生实体：属性包括学号、姓名、性别、出生日期、所在院系，其中学号为主码。')
    doc.add_paragraph('（2）课程实体：属性包括课程号、课程名、学分、先修课程号，其中课程号为主码。')
    doc.add_paragraph('（3）选课联系：学生与课程之间的多对多联系，属性包括成绩。')

    doc.add_heading('3.2 E-R图描述', level=2)
    doc.add_paragraph(
        '实体关系描述：\n'
        '  - 学生（Student）：学号(Sno)、姓名(Sname)、性别(Ssex)、出生日期(Sbirth)、院系(Sdept)\n'
        '  - 课程（Course）：课程号(Cno)、课程名(Cname)、学分(Credit)、先修课(Cpno)\n'
        '  - 选课（SC）：多对多联系，一个学生可选多门课程，一门课程可被多个学生选修\n'
        '    联系属性：成绩(Grade)'
    )
    doc.add_paragraph('【请在此处插入E-R图截图】', style='Intense Quote')
    doc.add_paragraph(
        'E-R图说明：\n'
        '  Student实体 —(选修)— SC联系 —(被选)— Course实体\n'
        '  Student:SC = 1:n (一个学生有多个选课记录)\n'
        '  Course:SC = 1:n (一门课程有多个选课记录)\n'
        '  实际上SC是Student和Course的多对多联系的分解'
    )

    # 四、逻辑结构设计
    doc.add_heading('四、逻辑结构设计', level=1)
    doc.add_paragraph('将E-R图转换为关系模式（MySQL数据库中的表结构）：')
    doc.add_paragraph(
        '学生表 student(Sno, Sname, Ssex, Sbirth, Sdept)\n'
        '  主码：Sno\n'
        '  约束：Ssex ∈ {男, 女}\n\n'
        '课程表 course(Cno, Cname, Credit, Cpno)\n'
        '  主码：Cno\n'
        '  外码：Cpno REFERENCES course(Cno)\n\n'
        '选课表 sc(Sno, Cno, Grade)\n'
        '  主码：(Sno, Cno)\n'
        '  外码：Sno REFERENCES student(Sno) ON DELETE CASCADE\n'
        '  外码：Cno REFERENCES course(Cno) ON DELETE CASCADE\n'
        '  约束：Grade ∈ [0, 100]'
    )
    doc.add_paragraph('【请在此处插入数据库表结构截图】', style='Intense Quote')

    # 五、物理结构设计
    doc.add_heading('五、物理结构设计（SQL实现）', level=1)

    doc.add_heading('5.1 创建数据库', level=2)
    doc.add_paragraph(
        'CREATE DATABASE `学生选课` CHARACTER SET utf8mb4 COLLATE utf8mb4_unicode_ci;'
    )
    doc.add_paragraph('【请在此处插入执行截图】', style='Intense Quote')

    doc.add_heading('5.2 创建表', level=2)
    doc.add_paragraph(
        '详见 SQL 脚本文件：实验十一_数据库设计.sql\n\n'
        '主要建表语句包括：\n'
        '- CREATE TABLE student (...)\n'
        '- CREATE TABLE course (...)\n'
        '- CREATE TABLE sc (...)\n'
        '其中student和course表设置了主键约束，sc表设置了复合主键和外键约束。'
    )
    doc.add_paragraph('【请在此处插入建表执行截图】', style='Intense Quote')

    doc.add_heading('5.3 插入数据', level=2)
    doc.add_paragraph(
        '向三张表中分别插入数据：\n'
        '- 7名学生记录（3个院系）\n'
        '- 8门课程记录（含先修课关系）\n'
        '- 16条选课记录（含成绩）'
    )
    doc.add_paragraph('【请在此处插入数据插入截图】', style='Intense Quote')

    # 六、Python数据库操作
    doc.add_heading('六、Python数据库操作', level=1)
    doc.add_paragraph('使用 pymysql 库连接MySQL数据库，实现以下功能：')

    doc.add_heading('6.1 连接数据库', level=2)
    doc.add_paragraph(
        'import pymysql\n'
        "connection = pymysql.connect(host='localhost', user='root',\n"
        "    password='********', charset='utf8mb4')\n"
        "cursor = connection.cursor()"
    )

    doc.add_heading('6.2 创建数据库和表', level=2)
    doc.add_paragraph(
        '通过Python执行CREATE DATABASE和CREATE TABLE语句，\n'
        '自动创建数据库和表结构。详见 mysql_operations.py'
    )

    doc.add_heading('6.3 插入数据', level=2)
    doc.add_paragraph(
        '使用参数化查询（%s占位符）插入数据，防止SQL注入。\n'
        '使用 INSERT IGNORE 避免重复插入。\n'
        '详见 mysql_operations.py 中的 insert_data() 函数。'
    )

    doc.add_heading('6.4 查询操作', level=2)
    doc.add_paragraph(
        '实现了以下查询：\n'
        '（1）查询所有学生信息\n'
        '（2）查询所有课程信息\n'
        '（3）查询学生选课情况及成绩（多表连接）\n'
        '（4）统计各院系学生人数（GROUP BY）\n'
        '（5）查询指定课程的学生成绩（条件查询）\n'
        '（6）交互式查询：根据输入查询最高成绩'
    )
    doc.add_paragraph('【请在此处插入Python执行结果截图】', style='Intense Quote')

    # 七、实验结果
    doc.add_heading('七、实验结果', level=1)

    doc.add_heading('7.1 数据表结构验证', level=2)
    doc.add_paragraph(
        'student表包含字段：Sno, Sname, Ssex, Sbirth, Sdept (7条记录)\n'
        'course表包含字段：Cno, Cname, Credit, Cpno (8条记录)\n'
        'sc表包含字段：Sno, Cno, Grade (16条记录)'
    )
    doc.add_paragraph('【请在此处插入DESCRIBE结果截图】', style='Intense Quote')

    doc.add_heading('7.2 查询结果统计', level=2)
    doc.add_paragraph(
        '- 学生总数：7人（计算机系4人，信息系2人，数学系1人）\n'
        '- 课程总数：8门\n'
        '- 选课记录：16条\n'
        '- 各课程平均成绩在75~90分之间\n'
        '- 数据库系统概论选课学生2人，平均成绩84分'
    )
    doc.add_paragraph('【请在此处插入查询结果截图】', style='Intense Quote')

    # 八、总结
    doc.add_heading('八、实验总结', level=1)
    doc.add_paragraph(
        '通过本次实验，我掌握了以下知识点：\n\n'
        '1. 数据库设计方法：从需求分析到概念结构设计（E-R图）、逻辑结构设计（关系模式）、'
        '物理结构设计（SQL建表），完整地实践了数据库设计的全过程。\n\n'
        '2. E-R模型：理解了实体、属性、联系的概念，掌握了E-R图向关系模式的转换规则，'
        '特别是多对多联系需要转换为独立的关系表（sc表）。\n\n'
        '3. SQL建表：掌握了CREATE TABLE语句的用法，包括主键约束、外键约束、'
        'CHECK约束的定义，以及ON DELETE CASCADE等引用完整性约束。\n\n'
        '4. Python数据库编程：学会了使用pymysql库连接MySQL数据库，'
        '执行SQL语句，进行数据的增删改查操作，以及使用参数化查询防止SQL注入。\n\n'
        '5. 数据库设计是一个系统工程，需要在概念设计阶段充分分析需求，'
        '在逻辑设计阶段合理设计关系模式，在物理设计阶段充分考虑性能优化。'
    )

    doc.save('E:/study/DatabaseHomework/ex11/实验十一_数据库设计实验报告.docx')
    print('实验十一 实验报告已生成！')


if __name__ == '__main__':
    create_report()
