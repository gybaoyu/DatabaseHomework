"""
生成实验十四 Word 文档
实验十四：事务及并发控制实验
"""
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime


def create_document():
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)

    # 标题
    title = doc.add_heading('实验十四 事务及并发控制实验', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('')
    info_lines = [
        '实验名称：事务及并发控制实验',
        f'实验日期：{datetime.date.today()}',
        '数据库：学生管理',
        '实验环境：MySQL 8.0 + Windows',
    ]
    for line in info_lines:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # 一、实验目的
    doc.add_heading('一、实验目的', level=1)
    doc.add_paragraph('1. 理解事务的基本概念，掌握COMMIT和ROLLBACK的用法。')
    doc.add_paragraph('2. 理解事务的ACID特性（原子性、一致性、隔离性、持久性）。')
    doc.add_paragraph('3. 理解MySQL事务隔离级别，特别是REPEATABLE READ和READ COMMITTED的区别。')
    doc.add_paragraph('4. 了解并发控制中可能出现的问题（脏读、不可重复读、幻读）。')

    # 二、COMMIT和ROLLBACK
    doc.add_heading('二、实验内容——COMMIT 和 ROLLBACK', level=1)

    doc.add_heading('2.1 查看实验前 course 表', level=2)
    doc.add_paragraph('SQL语句：\nSELECT * FROM course ORDER BY course_id;')
    doc.add_paragraph('【请在此处插入执行截图（实验前的course表数据）】', style='Intense Quote')

    doc.add_heading('2.2 ROLLBACK 演示', level=2)
    doc.add_paragraph(
        '在事务中插入3条课程记录，使用ROLLBACK撤销所有修改。\n\n'
        'SQL语句：\n'
        'START TRANSACTION;\n'
        "INSERT INTO course VALUES ('04010107','Python程序设计','考查',36,2.0,NULL);\n"
        "INSERT INTO course VALUES ('04010108','大数据技术','考试',40,2.5,NULL);\n"
        "INSERT INTO course VALUES ('04010109','人工智能导论','考查',48,3.0,NULL);\n"
        'SELECT * FROM course ORDER BY course_id;  -- 事务内看到3条新记录\n'
        'ROLLBACK;\n'
        'SELECT * FROM course ORDER BY course_id;  -- 回滚后3条记录消失\n'
    )
    doc.add_paragraph('【请在此处插入ROLLBACK前后的执行截图】', style='Intense Quote')
    doc.add_paragraph(
        '分析：ROLLBACK后，3条新插入的课程记录被撤销，course表恢复原状。'
        '这说明事务具有原子性——事务中的所有操作要么全部执行，要么全部不执行。'
    )

    doc.add_heading('2.3 COMMIT 演示', level=2)
    doc.add_paragraph(
        '在事务中插入3条课程记录，使用COMMIT永久保存。\n\n'
        'SQL语句：\n'
        'START TRANSACTION;\n'
        "INSERT INTO course VALUES ('04010107','Python程序设计','考查',36,2.0,NULL);\n"
        "INSERT INTO course VALUES ('04010108','大数据技术','考试',40,2.5,NULL);\n"
        "INSERT INTO course VALUES ('04010109','人工智能导论','考查',48,3.0,NULL);\n"
        'SELECT * FROM course ORDER BY course_id;  -- 事务内看到3条新记录\n'
        'COMMIT;\n'
        'SELECT * FROM course ORDER BY course_id;  -- 提交后3条记录被永久保存\n'
    )
    doc.add_paragraph('【请在此处插入COMMIT前后的执行截图】', style='Intense Quote')
    doc.add_paragraph(
        '分析：COMMIT后，3条新插入的课程记录被永久保存到数据库中。'
        '即使后续系统出现故障，这些数据也不会丢失，体现了事务的持久性（Durability）。'
    )

    # 三、事务隔离级别
    doc.add_heading('三、实验内容——事务隔离级别', level=1)

    doc.add_heading('3.1 查看当前事务隔离级别', level=2)
    doc.add_paragraph(
        'SQL语句：\n'
        'SELECT @@transaction_isolation;\n'
        'SELECT @@global.transaction_isolation;\n'
        'SELECT @@session.transaction_isolation;'
    )
    doc.add_paragraph('【请在此处插入执行截图（默认应为 REPEATABLE-READ）】', style='Intense Quote')

    doc.add_heading('3.2 修改隔离级别为 READ-COMMITTED', level=2)
    doc.add_paragraph(
        'SQL语句：\n'
        "SET SESSION TRANSACTION ISOLATION LEVEL READ COMMITTED;\n"
        'SELECT @@session.transaction_isolation;'
    )
    doc.add_paragraph('【请在此处插入执行截图】', style='Intense Quote')

    doc.add_heading('3.3 恢复隔离级别为 REPEATABLE-READ', level=2)
    doc.add_paragraph(
        'SQL语句：\n'
        "SET SESSION TRANSACTION ISOLATION LEVEL REPEATABLE READ;\n"
        'SELECT @@session.transaction_isolation;'
    )
    doc.add_paragraph('【请在此处插入执行截图】', style='Intense Quote')

    # 四、并发控制
    doc.add_heading('四、实验内容——并发控制', level=1)

    doc.add_heading('4.1 REPEATABLE READ 下的并发更新', level=2)
    doc.add_paragraph(
        '实验场景：两个会话T1和T2同时操作score表，T1查询，T2更新。\n\n'
        '步骤：\n'
        '（1）T1：START TRANSACTION; SELECT * FROM score WHERE student_id=\'2021094001\';\n'
        '（2）T2：START TRANSACTION;\n'
        '        UPDATE score SET score_grade=score_grade+2 WHERE student_id=\'2021094001\';\n'
        '        COMMIT;\n'
        '（3）T1：SELECT * FROM score WHERE student_id=\'2021094001\';\n'
        '        -- 在REPEATABLE READ下，T1看不到T2的修改\n'
        '（4）T1：COMMIT;\n'
        '        SELECT * FROM score WHERE student_id=\'2021094001\';\n'
        '        -- 提交后再次查询，可以看到T2的修改'
    )
    doc.add_paragraph('【请在此处插入T1和T2两窗口对照截图】', style='Intense Quote')
    doc.add_paragraph(
        '分析：在REPEATABLE READ隔离级别下，事务T1在整个事务过程中看到的数据是一致的，'
        '不会受到其他事务提交的影响（避免了不可重复读）。'
    )

    doc.add_heading('4.2 READ COMMITTED 下的并发更新', level=2)
    doc.add_paragraph(
        '实验场景：将T1的隔离级别设为READ COMMITTED，重复上述实验。\n\n'
        '步骤：\n'
        '（1）T1：SET SESSION TRANSACTION ISOLATION LEVEL READ COMMITTED;\n'
        '        START TRANSACTION;\n'
        '        SELECT * FROM score WHERE student_id=\'2021094001\';\n'
        '（2）T2：UPDATE score SET score_grade=score_grade+3 WHERE student_id=\'2021094001\';\n'
        '        COMMIT;\n'
        '（3）T1：SELECT * FROM score WHERE student_id=\'2021094001\';\n'
        '        -- 在READ COMMITTED下，T1能看到T2已提交的修改\n'
        '        COMMIT;'
    )
    doc.add_paragraph('【请在此处插入T1和T2两窗口对照截图】', style='Intense Quote')
    doc.add_paragraph(
        '分析：在READ COMMITTED隔离级别下，事务T1可以读到其他事务已提交的修改，'
        '因此T1两次查询的结果可能不同（存在不可重复读）。'
    )

    doc.add_heading('4.3 REPEATABLE READ 下的幻读测试', level=2)
    doc.add_paragraph(
        '实验场景：T1在事务中查询，T2插入新记录，观察T1是否能看到新插入的记录。\n\n'
        '步骤：\n'
        '（1）T1：START TRANSACTION;\n'
        '        SELECT * FROM score WHERE student_id=\'2021094001\';\n'
        '（2）T2：START TRANSACTION;\n'
        "        INSERT INTO score VALUES('2021094001','04010107',80,'202220231');\n"
        '        COMMIT;\n'
        '（3）T1：SELECT * FROM score WHERE student_id=\'2021094001\';\n'
        '        -- 在REPEATABLE READ下，T1看不到T2的插入（无幻读）\n'
        '        COMMIT;\n'
        '（4）T1：SELECT * FROM score WHERE student_id=\'2021094001\';\n'
        '        -- 提交后可以看到新插入的记录'
    )
    doc.add_paragraph('【请在此处插入执行截图】', style='Intense Quote')
    doc.add_paragraph(
        '分析：MySQL的InnoDB引擎在REPEATABLE READ隔离级别下通过MVCC（多版本并发控制）\n'
        '和间隙锁（Gap Lock）机制避免了幻读问题，这是MySQL与标准SQL的重要区别之一。'
    )

    # 五、实验总结
    doc.add_heading('五、实验总结', level=1)
    doc.add_paragraph(
        '通过本次实验，我掌握了以下知识点：\n\n'
        '1. 事务基本操作：\n'
        '   - START TRANSACTION 开始一个事务\n'
        '   - COMMIT 提交事务，使修改永久生效\n'
        '   - ROLLBACK 回滚事务，撤销所有未提交的修改\n\n'
        '2. 事务的ACID特性：\n'
        '   - 原子性（Atomicity）：事务的操作要么全做，要么全不做\n'
        '   - 一致性（Consistency）：事务使数据库从一个一致状态变到另一个一致状态\n'
        '   - 隔离性（Isolation）：并发事务之间互不干扰\n'
        '   - 持久性（Durability）：已提交的事务结果永久保存\n\n'
        '3. 事务隔离级别：\n'
        '   - READ UNCOMMITTED：最低级别，可能出现脏读、不可重复读、幻读\n'
        '   - READ COMMITTED：避免脏读，但可能出现不可重复读和幻读\n'
        '   - REPEATABLE READ（MySQL默认）：避免脏读和不可重复读\n'
        '     MySQL通过MVCC+间隙锁避免了幻读\n'
        '   - SERIALIZABLE：最高级别，完全串行化\n\n'
        '4. 并发控制是数据库系统中保证数据一致性的重要机制。'
    )

    doc.save('E:/study/DatabaseHomework/ex14/实验十四_事务及并发控制实验报告.docx')
    print('实验十四 Word 文档已生成！')


if __name__ == '__main__':
    create_document()
