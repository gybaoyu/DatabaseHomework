"""
生成实验十三 Word 文档
实验十三：自主存取控制实验
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import datetime


def create_document():
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)

    # ========== 封面/标题 ==========
    title = doc.add_heading('实验十三 自主存取控制实验', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('')
    info_lines = [
        '实验名称：自主存取控制（DAC）实验',
        f'实验日期：{datetime.date.today()}',
        '数据库：学生管理',
        '实验环境：MySQL 8.0 + Windows',
    ]
    for line in info_lines:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ========== 实验目的 ==========
    doc.add_heading('一、实验目的', level=1)
    doc.add_paragraph(
        '1. 掌握MySQL中用户管理的基本操作，包括创建用户、修改密码、删除用户。'
    )
    doc.add_paragraph(
        '2. 掌握MySQL中权限管理的方法，包括授予权限（GRANT）和收回权限（REVOKE）。'
    )
    doc.add_paragraph(
        '3. 理解MySQL中角色的概念，掌握角色的创建、权限授予、角色分配和收回。'
    )
    doc.add_paragraph(
        '4. 理解自主存取控制（DAC）的基本原理。'
    )

    # ========== 实验内容一：用户管理 ==========
    doc.add_heading('二、实验内容——用户管理', level=1)

    doc.add_heading('2.1 创建用户 teas1（密码 stu）和 t1（密码 t1）', level=2)
    doc.add_paragraph(
        '使用 CREATE USER 语句创建两个新用户：teas1 和 t1。'
    )
    doc.add_paragraph(
        'SQL语句：\n'
        "CREATE USER 'teas1'@'localhost' IDENTIFIED BY 'stu';\n"
        "CREATE USER 't1'@'localhost' IDENTIFIED BY 't1';"
    )
    doc.add_paragraph('【请在此处插入执行截图】', style='Intense Quote')

    doc.add_heading('2.2 修改用户 t1 的密码为 hello', level=2)
    doc.add_paragraph(
        '使用 ALTER USER 语句修改 t1 用户的密码。\n\n'
        'SQL语句：\n'
        "ALTER USER 't1'@'localhost' IDENTIFIED BY 'hello';"
    )
    doc.add_paragraph('【请在此处插入执行截图】', style='Intense Quote')

    doc.add_heading('2.3 删除用户 stu', level=2)
    doc.add_paragraph(
        '使用 DROP USER 语句删除用户 stu（如果存在）。\n\n'
        'SQL语句：\n'
        "DROP USER IF EXISTS 'stu'@'localhost';"
    )
    doc.add_paragraph('【请在此处插入执行截图】', style='Intense Quote')

    doc.add_heading('2.4 查看用户', level=2)
    doc.add_paragraph(
        '通过查询 mysql.user 表查看已创建的用户。\n\n'
        'SQL语句：\n'
        "SELECT User, Host FROM mysql.user WHERE User IN ('teas1', 't1', 'stu');"
    )
    doc.add_paragraph('【请在此处插入执行截图（显示 teas1 和 t1 用户）】', style='Intense Quote')

    # ========== 实验内容二：权限管理 ==========
    doc.add_heading('三、实验内容——权限管理', level=1)

    doc.add_heading('3.1 授予 t1 对 student 表的查询权限', level=2)
    doc.add_paragraph(
        'SQL语句：\n'
        "GRANT SELECT ON `学生管理`.student TO 't1'@'localhost';"
    )
    doc.add_paragraph('【请在此处插入执行截图】', style='Intense Quote')

    doc.add_heading('3.2 授予 t1 对 score 表的插入权限', level=2)
    doc.add_paragraph(
        'SQL语句：\n'
        "GRANT INSERT ON `学生管理`.score TO 't1'@'localhost';"
    )
    doc.add_paragraph('【请在此处插入执行截图】', style='Intense Quote')

    doc.add_heading('3.3 授予 t1 对 score 表的更新权限', level=2)
    doc.add_paragraph(
        'SQL语句：\n'
        "GRANT UPDATE ON `学生管理`.score TO 't1'@'localhost';"
    )
    doc.add_paragraph('【请在此处插入执行截图】', style='Intense Quote')

    doc.add_heading('3.4 授予 t1 对 score 和 student 表的列级查询权限', level=2)
    doc.add_paragraph(
        '授予 t1 对 score 表特定列（student_id, course_id, score_grade, score_semester）\n'
        '和 student 表特定列（student_id, student_name, student_sex）的 SELECT 权限。\n\n'
        'SQL语句：\n'
        'GRANT SELECT (student_id, course_id, score_grade, score_semester)\n'
        '    ON `学生管理`.score TO \'t1\'@\'localhost\';\n'
        'GRANT SELECT (student_id, student_name, student_sex)\n'
        '    ON `学生管理`.student TO \'t1\'@\'localhost\';'
    )
    doc.add_paragraph('【请在此处插入执行截图】', style='Intense Quote')

    doc.add_heading('3.5 查看 t1 的权限', level=2)
    doc.add_paragraph(
        '使用 SHOW GRANTS 语句查看 t1 用户当前拥有的所有权限。\n\n'
        'SQL语句：\n'
        "SHOW GRANTS FOR 't1'@'localhost';"
    )
    doc.add_paragraph('【请在此处插入执行截图（显示 t1 的所有权限）】', style='Intense Quote')

    doc.add_heading('3.6 收回 t1 对 score 表的查询权限', level=2)
    doc.add_paragraph(
        '使用 REVOKE 语句收回 t1 对 score 表的 SELECT 权限。\n\n'
        'SQL语句：\n'
        "REVOKE SELECT ON `学生管理`.score FROM 't1'@'localhost';"
    )
    doc.add_paragraph('【请在此处插入执行截图】', style='Intense Quote')

    doc.add_heading('3.7 创建视图 CS_Student 并授予 t1 查询权限', level=2)
    doc.add_paragraph(
        '创建 CS_Student 视图（信息系学生），并授予 t1 对该视图的查询权限。\n\n'
        'SQL语句：\n'
        'CREATE VIEW CS_Student AS\n'
        '    SELECT * FROM student WHERE class_id IN (\n'
        "        SELECT class_id FROM class WHERE class_major LIKE '%信息%'\n"
        '    );\n'
        "GRANT SELECT ON `学生管理`.CS_Student TO 't1'@'localhost';"
    )
    doc.add_paragraph('【请在此处插入执行截图】', style='Intense Quote')

    # ========== 实验内容三：角色管理 ==========
    doc.add_heading('四、实验内容——角色管理', level=1)

    doc.add_heading('4.1 创建角色 teacher', level=2)
    doc.add_paragraph(
        'SQL语句：\n'
        "CREATE ROLE 'teacher';"
    )
    doc.add_paragraph('【请在此处插入执行截图】', style='Intense Quote')

    doc.add_heading('4.2 授予 teacher 角色对 course 表的权限', level=2)
    doc.add_paragraph(
        '授予 teacher 角色对 course 表的 SELECT、INSERT、UPDATE、DELETE 权限。\n\n'
        'SQL语句：\n'
        "GRANT SELECT ON `学生管理`.course TO 'teacher';\n"
        "GRANT INSERT, UPDATE, DELETE ON `学生管理`.course TO 'teacher';"
    )
    doc.add_paragraph('【请在此处插入执行截图】', style='Intense Quote')

    doc.add_heading('4.3 查看 teacher 角色权限', level=2)
    doc.add_paragraph(
        'SQL语句：\n'
        "SHOW GRANTS FOR 'teacher';"
    )
    doc.add_paragraph('【请在此处插入执行截图（显示 teacher 角色的权限）】', style='Intense Quote')

    doc.add_heading('4.4 将 teacher 角色授予 t1 用户', level=2)
    doc.add_paragraph(
        '使用 GRANT 语句将 teacher 角色分配给 t1 用户，使 t1 继承 teacher 角色的所有权限。\n\n'
        'SQL语句：\n'
        "GRANT 'teacher' TO 't1'@'localhost';"
    )
    doc.add_paragraph('【请在此处插入执行截图】', style='Intense Quote')

    doc.add_heading('4.5 收回 teacher 角色的 DELETE 权限', level=2)
    doc.add_paragraph(
        '使用 REVOKE 语句收回 teacher 角色对 course 表的 DELETE 权限。\n\n'
        'SQL语句：\n'
        "REVOKE DELETE ON `学生管理`.course FROM 'teacher';"
    )
    doc.add_paragraph('【请在此处插入执行截图】', style='Intense Quote')

    doc.add_heading('4.6 从 t1 收回 teacher 角色', level=2)
    doc.add_paragraph(
        'SQL语句：\n'
        "REVOKE 'teacher' FROM 't1'@'localhost';"
    )
    doc.add_paragraph('【请在此处插入执行截图】', style='Intense Quote')

    doc.add_heading('4.7 查看收回角色后 t1 的权限', level=2)
    doc.add_paragraph(
        '查看 t1 当前权限，确认 teacher 角色已被收回。\n\n'
        'SQL语句：\n'
        "SHOW GRANTS FOR 't1'@'localhost';"
    )
    doc.add_paragraph('【请在此处插入执行截图（t1 已无 teacher 角色）】', style='Intense Quote')

    # ========== 实验总结 ==========
    doc.add_heading('五、实验总结', level=1)
    doc.add_paragraph(
        '通过本次实验，我掌握了以下知识点：\n\n'
        '1. 用户管理：学会使用 CREATE USER、ALTER USER、DROP USER 等语句管理MySQL用户。\n\n'
        '2. 权限管理：理解了GRANT和REVOKE语句的用法，包括表级权限和列级权限的授予与收回。\n'
        '   - GRANT授予权限，REVOKE收回权限\n'
        '   - 可以授予表级权限（SELECT ON table）和列级权限（SELECT(col1, col2) ON table）\n\n'
        '3. 角色管理：理解了角色的概念及其在权限管理中的作用。\n'
        '   - CREATE ROLE 创建角色\n'
        '   - GRANT ... TO role 给角色授权\n'
        '   - GRANT role TO user 将角色授予用户\n'
        '   - REVOKE ... FROM role/user 收回权限或角色\n\n'
        '4. 自主存取控制（DAC）允许数据库对象的创建者自主决定谁可以访问这些对象，\n'
        '   以及授予何种访问权限，这是数据库安全的重要机制之一。'
    )

    # 保存文档
    doc.save('E:/study/DatabaseHomework/ex13/实验十三_自主存取控制实验报告.docx')
    print('实验十三 Word 文档已生成！')


if __name__ == '__main__':
    create_document()
